import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def write_file(text_list, filename):
    new_doc = docx.Document() # Create output docx

    title = new_doc.add_paragraph('')
    title.add_run(f'"{filename}":\n').bold = True # Make the title bold
    title.runs[0].font.size = docx.shared.Pt(20) # Change the font size to 20
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # Center align the paragraph
    for run_list in text_list:
        para = new_doc.add_paragraph('')
        for text in run_list:
            if isinstance(text, list):
                para.add_run('\n' + text[0] + ': ').bold = True
            else:
                para.add_run(text)
    print(f'\nSaved file as: {filename}')
    new_doc.save(filename)

if __name__ == "__main__":
    pass
